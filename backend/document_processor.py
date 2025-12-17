"""
Multi-Format Document Processor (PDF, DOCX, TXT)

Supports:
- PDF: Smart OCR detection (Nanonets for scanned, direct extraction for searchable)
- DOCX: Direct text extraction with formatting preservation
- TXT: Plain text with basic cleaning

All formats output clean, chunked documents ready for RAG.
"""

import fitz  # PyMuPDF
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from config import settings
import re
import os
import requests
import time
import logging

# Configure logger
logging.basicConfig(
    level=logging.INFO,
    format='%(message)s',
    force=True
)
logger = logging.getLogger(__name__)

# Nanonets API configuration (for PDF OCR)
NANONETS_API_KEY = os.getenv("NANONETS_API_KEY", "7d07e046-8665-11f0-a771-3ea31997c197")
NANONETS_ASYNC_URL = "https://extraction-api.nanonets.com/api/v1/extract/async"
NANONETS_RESULTS_URL = "https://extraction-api.nanonets.com/api/v1/extract/results"

# Polling configuration
DEFAULT_POLLING_INTERVAL = 5.0
DEFAULT_MAX_WAIT_TIME = 180.0


def clean_markdown_text(text: str) -> str:
    """Clean markdown text for better RAG performance."""
    if not text:
        return ""
    
    # Decode HTML entities
    text = text.replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&amp;', '&').replace('&quot;', '"')
    
    # Remove image tags
    text = re.sub(r'<img[^>]*>[^<]*</img>', '', text)
    text = re.sub(r'<img[^>]*/?>', '', text)
    
    # Remove excessive markdown formatting
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Clean up table markup
    text = re.sub(r'</?table[^>]*>', '\n', text)
    text = re.sub(r'</?thead[^>]*>', '', text)
    text = re.sub(r'</?tbody[^>]*>', '', text)
    text = re.sub(r'</?tr[^>]*>', '\n', text)
    text = re.sub(r'</?th[^>]*>', ' | ', text)
    text = re.sub(r'</?td[^>]*>', ' | ', text)
    
    # Clean up whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\| +\|', '|', text)
    
    return text.strip()


def calculate_text_quality(text: str) -> Tuple[float, Dict[str, Any]]:
    """Calculate quality score for extracted text (0.0 to 1.0)."""
    if not text or len(text) < 100:
        return 0.0, {"reason": "Text too short"}
    
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    word_count = len(words)
    
    if word_count < 20:
        return 0.0, {"reason": "Too few words", "word_count": word_count}
    
    # Medical terminology
    medical_terms = ['patient', 'diagnosis', 'treatment', 'medical', 'doctor', 
                     'hospital', 'clinical', 'procedure', 'medication', 'symptoms']
    medical_count = sum(1 for term in medical_terms if term in text.lower())
    
    # Special characters ratio
    special_chars = len(re.findall(r'[^a-zA-Z0-9\s.,;:!?()-]', text))
    special_ratio = special_chars / len(text) if len(text) > 0 else 1.0
    
    # Sentence structure
    sentences = re.split(r'[.!?]+', text)
    avg_sentence_length = sum(len(s.split()) for s in sentences) / max(len(sentences), 1)
    
    # Calculate score
    word_score = min(word_count / 100, 1.0) * 0.4
    medical_score = min(medical_count / 3, 1.0) * 0.3
    special_penalty = max(0, 1.0 - special_ratio * 2) * 0.2
    sentence_score = min(avg_sentence_length / 15, 1.0) * 0.1
    
    quality_score = word_score + medical_score + special_penalty + sentence_score
    
    details = {
        "word_count": word_count,
        "medical_terms": medical_count,
        "special_ratio": round(special_ratio, 3),
        "avg_sentence_length": round(avg_sentence_length, 1)
    }
    
    return quality_score, details


class DocumentProcessor:
    """
    Multi-format document processor supporting PDF, DOCX, and TXT.
    
    Features:
    - PDF: Smart OCR detection + Nanonets integration
    - DOCX: Direct text extraction with paragraph structure
    - TXT: Plain text with basic cleaning
    """
    
    def __init__(self):
        self.text_splitter = SentenceSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP
        )
    
    def _smart_chunk(self, text: str) -> List[str]:
        """
        Smart chunking that adapts based on document size.
        
        For small documents (< 3000 chars / ~1 page), returns the entire text as one chunk
        to prevent losing content at chunk boundaries. Most 1-page clinical notes are 1500-2500 chars.
        
        For larger documents, uses normal sentence-based chunking with 40% overlap.
        """
        # Small documents (1-page clinical notes): keep as single chunk to preserve all content
        if len(text) < 3000:
            logger.info(f"   📄 Small document ({len(text)} chars) - indexing as single chunk")
            return [text]
        
        # Large documents (2+ pages): use normal chunking with overlap
        chunks = self.text_splitter.split_text(text)
        logger.info(f"   📄 Large document ({len(text)} chars) - split into {len(chunks)} chunks")
        return chunks
    
    def create_documents(self, file_path: str) -> List[Document]:
        """
        Process any supported document format and return chunked documents.
        
        Args:
            file_path: Path to document (PDF, DOCX, or TXT)
        
        Returns:
            List of LlamaIndex Document objects with metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Route to appropriate processor based on extension
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension == '.docx':
            return self._process_docx(file_path)
        elif extension == '.txt':
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}. Supported: .pdf, .docx, .txt")
    
    def _process_txt(self, file_path: Path) -> List[Document]:
        """Process plain text file."""
        logger.info(f"\n{'='*60}")
        logger.info(f"📄 Processing TXT: {file_path.name}")
        logger.info(f"{'='*60}")
        
        try:
            # Read text file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Basic cleaning
            text = clean_markdown_text(text)
            
            if not text or len(text) < 50:
                logger.warning("⚠️  Text file is empty or too short")
                return []
            
            # Calculate quality
            quality_score, quality_details = calculate_text_quality(text)
            
            # Chunk the text (smart chunking for small docs)
            chunks = self._smart_chunk(text)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    text=chunk,
                    metadata={
                        "source": file_path.name,
                        "file_type": "txt",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "ocr_method": "direct",
                        "ocr_quality": quality_score
                    }
                )
                documents.append(doc)
            
            logger.info(f"✅ TXT processed successfully")
            logger.info(f"   Quality: {quality_score:.2f}")
            logger.info(f"   Chunks: {len(documents)}")
            logger.info(f"   Total text: {len(text)} chars")
            logger.info(f"{'='*60}\n")
            
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error processing TXT: {e}")
            return []
    
    def _process_docx(self, file_path: Path) -> List[Document]:
        """Process DOCX file."""
        logger.info(f"\n{'='*60}")
        logger.info(f"📄 Processing DOCX: {file_path.name}")
        logger.info(f"{'='*60}")
        
        try:
            from docx import Document as DocxDocument
            
            # Read DOCX
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        paragraphs.append(row_text)
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            
            # Clean text
            full_text = clean_markdown_text(full_text)
            
            if not full_text or len(full_text) < 50:
                logger.warning("⚠️  DOCX file is empty or too short")
                return []
            
            # Calculate quality
            quality_score, quality_details = calculate_text_quality(full_text)
            
            # Chunk the text (smart chunking for small docs)
            chunks = self._smart_chunk(full_text)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    text=chunk,
                    metadata={
                        "source": file_path.name,
                        "file_type": "docx",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "ocr_method": "direct",
                        "ocr_quality": quality_score,
                        "paragraph_count": len(paragraphs)
                    }
                )
                documents.append(doc)
            
            logger.info(f"✅ DOCX processed successfully")
            logger.info(f"   Quality: {quality_score:.2f}")
            logger.info(f"   Paragraphs: {len(paragraphs)}")
            logger.info(f"   Chunks: {len(documents)}")
            logger.info(f"   Total text: {len(full_text)} chars")
            logger.info(f"{'='*60}\n")
            
            return documents
            
        except ImportError:
            logger.error("❌ python-docx not installed. Run: pip install python-docx")
            return []
        except Exception as e:
            logger.error(f"❌ Error processing DOCX: {e}")
            return []
    
    def _process_pdf(self, file_path: Path) -> List[Document]:
        """Process PDF with smart OCR detection (existing logic)."""
        logger.info(f"\n{'='*60}")
        logger.info(f"📄 Processing PDF: {file_path.name}")
        logger.info(f"{'='*60}")
        
        # Check if PDF is searchable
        is_searchable, page_count = self._is_pdf_searchable(str(file_path))
        
        if is_searchable:
            logger.info("✅ PDF is searchable - using direct extraction")
            return self._extract_text_direct(file_path)
        else:
            logger.info("🔍 PDF is scanned - using OCR")
            return self._extract_text_with_ocr(file_path)
    
    def _is_pdf_searchable(self, pdf_path: str) -> Tuple[bool, int]:
        """Check if PDF has extractable text."""
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # Sample first 3 pages
            sample_pages = min(3, page_count)
            total_text_length = 0
            
            for page_num in range(sample_pages):
                page = doc[page_num]
                text = page.get_text()
                total_text_length += len(text.strip())
            
            doc.close()
            
            # If we have reasonable text, it's searchable
            avg_text_per_page = total_text_length / sample_pages
            is_searchable = avg_text_per_page > 100
            
            return is_searchable, page_count
            
        except Exception as e:
            logger.error(f"Error checking PDF: {e}")
            return False, 0
    
    def _extract_text_direct(self, file_path: Path) -> List[Document]:
        """Extract text directly from searchable PDF."""
        try:
            doc = fitz.open(str(file_path))
            full_text = ""
            page_texts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                page_texts.append((page_num + 1, text))
                full_text += f"\n\n{text}"
            
            doc.close()
            
            # Clean text
            full_text = clean_markdown_text(full_text)
            
            # Calculate quality
            quality_score, _ = calculate_text_quality(full_text)
            
            # Chunk text (smart chunking for small docs)
            chunks = self._smart_chunk(full_text)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    text=chunk,
                    metadata={
                        "source": file_path.name,
                        "file_type": "pdf",
                        "page_number": f"1-{len(page_texts)}",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "ocr_method": "direct",
                        "ocr_quality": quality_score
                    }
                )
                documents.append(doc)
            
            logger.info(f"✅ PDF processed (direct extraction)")
            logger.info(f"   Quality: {quality_score:.2f}")
            logger.info(f"   Pages: {len(page_texts)}")
            logger.info(f"   Chunks: {len(documents)}")
            logger.info(f"{'='*60}\n")
            
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error in direct extraction: {e}")
            return []
    
    def _extract_text_with_ocr(self, file_path: Path) -> List[Document]:
        """Extract text from scanned PDF using Nanonets OCR."""
        if not NANONETS_API_KEY:
            logger.warning("⚠️  NANONETS_API_KEY not set, skipping OCR")
            return []
        
        try:
            # Submit to Nanonets
            logger.info("📤 Submitting to Nanonets OCR...")
            
            with open(file_path, 'rb') as f:
                response = requests.post(
                    NANONETS_ASYNC_URL,
                    auth=requests.auth.HTTPBasicAuth(NANONETS_API_KEY, ''),
                    files={'file': f}
                )
            
            if response.status_code != 200:
                logger.error(f"❌ Nanonets submission failed: {response.status_code}")
                return []
            
            result = response.json()
            record_id = result.get('record_id')
            
            if not record_id:
                logger.error("❌ No record_id returned from Nanonets")
                return []
            
            logger.info(f"✅ Submitted to Nanonets (record_id: {record_id})")
            logger.info("⏳ Waiting for OCR processing...")
            
            # Poll for results
            start_time = time.time()
            while time.time() - start_time < DEFAULT_MAX_WAIT_TIME:
                time.sleep(DEFAULT_POLLING_INTERVAL)
                
                status_response = requests.get(
                    f"{NANONETS_RESULTS_URL}/{record_id}",
                    auth=requests.auth.HTTPBasicAuth(NANONETS_API_KEY, '')
                )
                
                if status_response.status_code != 200:
                    continue
                
                status_data = status_response.json()
                status = status_data.get('status')
                
                if status == 'COMPLETE':
                    logger.info("✅ OCR complete!")
                    
                    # Extract markdown text
                    markdown_text = status_data.get('markdown', '')
                    
                    if not markdown_text:
                        logger.error("❌ No markdown text in OCR results")
                        return []
                    
                    # Clean and chunk
                    markdown_text = clean_markdown_text(markdown_text)
                    quality_score, _ = calculate_text_quality(markdown_text)
                    
                    # Chunk text (smart chunking for small docs)
                    chunks = self._smart_chunk(markdown_text)
                    
                    # Create documents
                    documents = []
                    for i, chunk in enumerate(chunks):
                        doc = Document(
                            text=chunk,
                            metadata={
                                "source": file_path.name,
                                "file_type": "pdf",
                                "page_number": "OCR",
                                "chunk_index": i,
                                "total_chunks": len(chunks),
                                "ocr_method": "nanonets",
                                "ocr_quality": quality_score
                            }
                        )
                        documents.append(doc)
                    
                    logger.info(f"✅ PDF processed (Nanonets OCR)")
                    logger.info(f"   Quality: {quality_score:.2f}")
                    logger.info(f"   Chunks: {len(documents)}")
                    logger.info(f"{'='*60}\n")
                    
                    return documents
                
                elif status == 'FAILED':
                    logger.error("❌ Nanonets OCR failed")
                    return []
            
            logger.error("❌ OCR timeout")
            return []
            
        except Exception as e:
            logger.error(f"❌ Error in OCR: {e}")
            return []


# Backwards compatibility
PDFProcessor = DocumentProcessor


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        processor = DocumentProcessor()
        docs = processor.create_documents(sys.argv[1])
        print(f"\n✅ Created {len(docs)} document chunks")
    else:
        print("Usage: python document_processor.py <path_to_file>")

